"""
Experiment document export
==========================
Finds the Observation Table of the experiment document and inserts a
"Result & Graph" section immediately below it containing:

  - the readings actually measured in the app (a real Word table), and
  - the chart rendered in the browser (PNG).

The experiment document is bundled with the backend (see assets/), so the
student just clicks "Add to Document" and downloads the finished file — no
upload step. The bundled bytes are read fresh on every request and never
modified; each export produces a brand-new document. An uploaded .docx is still
accepted (optional) for flexibility, but the app no longer requires one.
"""

import io
import json
import os
import sys
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse

router = APIRouter()

# Header text that identifies the Observation Table in the syllabus document.
_TABLE_MARKERS = ("reading no", "distance", "signal strength")

# Experiment key -> bundled template filename (in assets/). The template is the
# blank syllabus document; the export inserts the Result & Graph section below
# its Observation Table, so students never upload anything.
_TEMPLATES = {
    "exp4": "Expt. No. 4.docx",
    "exp5": "Expt No. 5.docx",
    "exp6": "Expt. No. 6.docx",
}


def _assets_dir() -> str:
    """
    Directory holding the bundled template documents.

    Resolves both when running from source and inside the PyInstaller one-file
    exe, where data files are unpacked to sys._MEIPASS at runtime.
    """
    if getattr(sys, "frozen", False):
        return os.path.join(sys._MEIPASS, "assets")
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")


def _load_template(key: str) -> bytes:
    """Reads a bundled template's bytes, or raises a clear HTTP error."""
    fname = _TEMPLATES.get(key)
    if not fname:
        raise HTTPException(status_code=400, detail=f"Unknown experiment template '{key}'.")
    path = os.path.join(_assets_dir(), fname)
    try:
        with open(path, "rb") as fh:
            return fh.read()
    except FileNotFoundError:
        raise HTTPException(
            status_code=500,
            detail="The bundled experiment document is missing from this build of the backend.",
        )


def _find_observation_table(doc: Document):
    """
    Locates the Observation Table. Matches on header text rather than position
    so it still works if the student's copy has extra tables or reordered
    sections; falls back to the first table in the document.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        header = " | ".join(c.text.strip().lower() for c in table.rows[0].cells)
        if sum(m in header for m in _TABLE_MARKERS) >= 2:
            return table
    return doc.tables[0] if doc.tables else None


def _style_header_cell(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)


def _apply_grid_borders(table) -> None:
    """
    Draws borders via direct XML rather than a named style.

    Not every .docx defines the built-in "Table Grid" style — python-docx raises
    KeyError if it is missing, and student copies of the syllabus document vary.
    Writing w:tblBorders directly always works, whatever styles the document has.
    """
    tblPr = table._element.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tblPr.append(borders)


def _quality(rssi) -> str:
    """
    Labels a measured RSSI, using the same thresholds as the signal badges in
    the web app so the document matches what the student saw on screen.

    Derived here because a reading only carries the measured values (rssi,
    signal_pct, ssid, channel) — there is no 'quality' field to read.
    """
    try:
        v = float(rssi)
    except (TypeError, ValueError):
        return ""
    if v >= -50:
        return "Excellent"
    if v >= -60:
        return "Good"
    if v >= -70:
        return "Fair"
    return "Poor"


def _num(value) -> str:
    """Formats a measured number without a pointless trailing '.0'."""
    try:
        v = float(value)
    except (TypeError, ValueError):
        return "" if value is None else str(value)
    return str(int(v)) if v == int(v) else f"{v:g}"


def _build_readings_table(doc: Document, readings: list, style_hint=None) -> "object":
    """Creates (at the end of the body) a Word table of the measured readings."""
    cols = ["Reading No.", "Distance (m)", "RSSI (dBm)", "Signal (%)", "Quality"]
    table = doc.add_table(rows=1, cols=len(cols))

    # Prefer the document's own Observation-Table style so the inserted table
    # looks native; fall back to Table Grid; borders are drawn regardless.
    for candidate in (style_hint, "Table Grid"):
        if not candidate:
            continue
        try:
            table.style = candidate
            break
        except (KeyError, ValueError):
            continue
    _apply_grid_borders(table)

    hdr = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr[i].text = name
        _style_header_cell(hdr[i])

    for idx, r in enumerate(readings, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = _num(r.get("distance"))
        row[2].text = _num(r.get("rssi"))
        row[3].text = _num(r.get("signal_pct"))
        row[4].text = str(r.get("quality") or _quality(r.get("rssi")))
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def _summary_sentence(readings: list) -> str:
    vals = [r.get("rssi") for r in readings if isinstance(r.get("rssi"), (int, float))]
    dists = [r.get("distance") for r in readings if isinstance(r.get("distance"), (int, float))]
    if not vals or not dists:
        return "Measured Wi-Fi signal strength at multiple distances from the access point."
    return (
        f"Across {len(readings)} readings taken between {min(dists):g} m and {max(dists):g} m, "
        f"the measured RSSI ranged from {max(vals):g} dBm (strongest) to {min(vals):g} dBm "
        f"(weakest), a total variation of {abs(max(vals) - min(vals)):g} dB. This confirms the "
        f"inverse, non-linear relationship between received signal strength and distance "
        f"described by the log-distance path loss model."
    )


# Blank paragraphs between the readings table and the graph. The syllabus
# document keeps the graph lower on the page (where the "Paste the Graph"
# placeholder used to sit), so we reproduce that spacing.
_GRAPH_GAP = 11

# Placeholder in the template that told students to paste their graph by hand.
# The export inserts the real graph, so this line (and the blank lines around it)
# are removed.
_GRAPH_PLACEHOLDER = "paste the graph"


def _p_is_empty(el) -> bool:
    return el.tag.endswith("}p") and not "".join(el.itertext()).strip()


def _remove_graph_placeholder(doc: Document) -> None:
    """
    Deletes the 'Paste the Graph' placeholder paragraph and one blank paragraph
    on each side of it, so the auto-inserted graph replaces it cleanly.
    """
    for para in list(doc.paragraphs):
        if para.text.strip().lower() != _GRAPH_PLACEHOLDER:
            continue
        el = para._element
        prev, nxt = el.getprevious(), el.getnext()
        if prev is not None and _p_is_empty(prev):
            prev.getparent().remove(prev)
        if nxt is not None and _p_is_empty(nxt):
            nxt.getparent().remove(nxt)
        el.getparent().remove(el)
        return


# ────────────────────────────────────────────────────────────────────────────
# Experiment 5 — Network Throughput & Latency
#
# The Exp-5 syllabus document has two observation tables (Table 1: CLI ping,
# Table 2: online throughput). The measured "Result" section is inserted below
# the throughput table (the last observation table), before the Conclusion.
# ────────────────────────────────────────────────────────────────────────────

# Header markers that identify the throughput table (Table 2) — the anchor.
_EXP5_ANCHOR_MARKERS = ("platform", "download throughput", "upload throughput")


def _find_exp5_anchor(doc: Document):
    """
    Locates the throughput table (Table 2) so the Result section can be inserted
    directly beneath it. Matches on header text; falls back to the last table so
    the result still lands below the observation section if headers were edited.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        header = " | ".join(c.text.strip().lower() for c in table.rows[0].cells)
        if sum(m in header for m in _EXP5_ANCHOR_MARKERS) >= 2:
            return table
    return doc.tables[-1] if doc.tables else None


def _apply_table_style(table, style_hint=None) -> None:
    """Applies the document's own table style if available, then draws borders."""
    for candidate in (style_hint, "Table Grid"):
        if not candidate:
            continue
        try:
            table.style = candidate
            break
        except (KeyError, ValueError):
            continue
    _apply_grid_borders(table)


def _fill_row(cells, values) -> None:
    """Writes values into a row and shrinks the font to match the sheet tables."""
    for cell, value in zip(cells, values):
        cell.text = "" if value is None else str(value)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)


def _build_cli_table(doc: Document, ping: dict, style_hint=None):
    """Table 1 — the student's real command-line ping measurement."""
    cols = [
        "Target Host", "Packets Sent", "Packets Received", "Packet Loss (%)",
        "Minimum Latency (ms)", "Maximum Latency (ms)", "Average Latency (ms)",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    _apply_table_style(table, style_hint)
    hdr = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr[i].text = name
        _style_header_cell(hdr[i])
    _fill_row(table.add_row().cells, [
        ping.get("host", ""),
        _num(ping.get("sent")),
        _num(ping.get("received")),
        _num(ping.get("packet_loss")),
        _num(ping.get("min_rtt")),
        _num(ping.get("max_rtt")),
        _num(ping.get("avg_rtt")),
    ])
    return table


def _build_throughput_table(doc: Document, speed: dict, platform: str, style_hint=None):
    """Table 2 — the student's real online speed-test measurement."""
    cols = [
        "Platform", "Download Throughput (Mbps)",
        "Upload Throughput (Mbps)", "Latency / Ping (ms)",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    _apply_table_style(table, style_hint)
    hdr = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr[i].text = name
        _style_header_cell(hdr[i])
    _fill_row(table.add_row().cells, [
        platform,
        _num(speed.get("download_mbps")),
        _num(speed.get("upload_mbps")),
        _num(speed.get("ping_ms")),
    ])
    return table


def _exp5_summary(ping: Optional[dict], speed: Optional[dict]) -> str:
    """One-line observation summarising whichever measurements are present."""
    parts = []
    if speed:
        parts.append(
            f"The live speed test measured {_num(speed.get('download_mbps'))} Mbps download "
            f"and {_num(speed.get('upload_mbps'))} Mbps upload throughput at a ping of "
            f"{_num(speed.get('ping_ms'))} ms (jitter {_num(speed.get('jitter_ms'))} ms)."
        )
    if ping:
        parts.append(
            f"The command-line ping to {ping.get('host', 'the target host')} sent "
            f"{_num(ping.get('sent'))} packets with {_num(ping.get('packet_loss'))}% loss and an "
            f"average round-trip time of {_num(ping.get('avg_rtt'))} ms "
            f"(min {_num(ping.get('min_rtt'))} ms, max {_num(ping.get('max_rtt'))} ms)."
        )
    if not parts:
        return "Measured the network's throughput and latency using real tools."
    return " ".join(parts) + (
        " Together these confirm that throughput and latency are independent pillars of "
        "network quality."
    )


def _sublabel(doc: Document, text: str):
    """A bold caption above an inserted result table."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    return p


# ─────────────────────── Experiment 6 (Bluetooth) ───────────────────────

# The last observation table of the Exp-6 document (Signal Quality vs Distance).
_EXP6_ANCHOR_MARKERS = ("paced distance", "link connectivity", "performance")


def _find_exp6_anchor(doc: Document):
    """Locate Table 2 (range field test) so the Result section lands beneath it."""
    for table in doc.tables:
        if not table.rows:
            continue
        header = " | ".join(c.text.strip().lower() for c in table.rows[0].cells)
        if sum(m in header for m in _EXP6_ANCHOR_MARKERS) >= 2:
            return table
    return doc.tables[-1] if doc.tables else None


def _bt_link_state(rssi):
    """
    Derive (connectivity state, stream-performance note) from a MEASURED BLE RSSI.

    Not a fabrication — it classifies the real dBm the app recorded, using the
    same boundaries as this experiment's reference simulator, and mirrors the
    wording of the syllabus Table 2 (Connected / Weak / Intermittent / Disconnected).
    """
    if rssi is None:
        return ("Disconnected", "No advertisement received — device out of range.")
    r = float(rssi)
    if r >= -65:
        return ("Connected", "Flawless connection. Instantaneous file/audio transfers.")
    if r >= -75:
        return ("Connected", "Very stable. Minimal buffering, stuttering, or delay.")
    if r >= -88:
        return ("Weak Connection", "Sluggish response times. Transfers take longer to start.")
    if r >= -98:
        return ("Intermittent", "Audio/data stream stutters, drops, and tries to reconnect.")
    return ("Disconnected", "Pairing link lost entirely due to signal path loss.")


def _bt_category(dev: dict) -> str:
    """Honest device category from the scan (BLE + vendor if the OUI resolved)."""
    vendor = (dev.get("vendor") or "").strip()
    return f"BLE · {vendor}" if vendor else "BLE (Low Energy)"


def _bt_services(dev: dict) -> str:
    """What service profiles were actually observed for this device."""
    if dev.get("connected") and dev.get("services_count"):
        n = dev["services_count"]
        return f"{n} GATT service{'s' if n != 1 else ''} (connected)"
    return "GAP advertisement only"


def _build_bt_discovery_table(doc: Document, devices: list, style_hint=None):
    """Table 1 — the devices actually discovered in the live BLE scan."""
    cols = ["Discovered Name", "MAC Address (BD_ADDR)", "Class / Category", "Supported Profile Services"]
    table = doc.add_table(rows=1, cols=len(cols))
    _apply_table_style(table, style_hint)
    hdr = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr[i].text = name
        _style_header_cell(hdr[i])
    for dev in devices:
        _fill_row(table.add_row().cells, [
            dev.get("name") or "(no name)",
            dev.get("address", ""),
            _bt_category(dev),
            _bt_services(dev),
        ])
    return table


def _build_bt_connected_table(doc: Document, conn: dict, style_hint=None):
    """The single device the student actually paired/connected to (Part B)."""
    cols = ["Connected Device", "MAC Address (BD_ADDR)", "Bond / Pairing State", "GATT Services Exposed"]
    table = doc.add_table(rows=1, cols=len(cols))
    _apply_table_style(table, style_hint)
    hdr = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr[i].text = name
        _style_header_cell(hdr[i])
    paired = conn.get("paired")
    bond = "Bonded (paired)" if paired else ("Connected (link only)" if conn.get("connected") else "Not connected")
    sc = conn.get("services_count")
    services = f"{sc} GATT service{'s' if sc != 1 else ''}" if sc else "—"
    _fill_row(table.add_row().cells, [
        conn.get("name") or "(no name)",
        conn.get("address", ""),
        bond,
        services,
    ])
    return table


def _build_bt_field_table(doc: Document, readings: list, style_hint=None):
    """Table 2 — measured RSSI at each paced distance, classified into link state."""
    cols = ["Paced Distance (m)", "Measured RSSI (dBm)", "Link Connectivity State", "Practical Data/Stream Performance"]
    table = doc.add_table(rows=1, cols=len(cols))
    _apply_table_style(table, style_hint)
    hdr = table.rows[0].cells
    for i, name in enumerate(cols):
        hdr[i].text = name
        _style_header_cell(hdr[i])
    for r in sorted(readings, key=lambda x: _as_float(x.get("distance"))):
        state, perf = _bt_link_state(r.get("rssi"))
        _fill_row(table.add_row().cells, [
            _num(r.get("distance")), _num(r.get("rssi")), state, perf,
        ])
    return table


def _as_float(v) -> float:
    try:
        return float(v)
    except (TypeError, ValueError):
        return 0.0


def _exp6_summary(devices: list, readings: list) -> str:
    parts = []
    if devices:
        named = [d for d in devices if d.get("name")]
        parts.append(
            f"The live BLE inquiry scan discovered {len(devices)} device"
            f"{'s' if len(devices) != 1 else ''} in piconet range"
            + (f", {len(named)} broadcasting a friendly name." if named else ".")
        )
    if readings:
        rssis = [_as_float(r.get("rssi")) for r in readings if r.get("rssi") is not None]
        dists = [_as_float(r.get("distance")) for r in readings if r.get("distance") is not None]
        if rssis and dists:
            parts.append(
                f"Across {len(readings)} paced distances from {min(dists):g} m to {max(dists):g} m, "
                f"the measured RSSI fell from {max(rssis):g} dBm to {min(rssis):g} dBm, confirming that "
                "Class 2 Bluetooth links stay stable within ~5–10 m before path loss drives the connection "
                "into stuttering and eventual link loss."
            )
    if not parts:
        return "Studied Bluetooth discovery, pairing and range using live device measurements."
    return " ".join(parts)


def _para_text(el) -> Optional[str]:
    """Plain text of a <w:p> body element, or None if it is not a paragraph."""
    if not el.tag.endswith("}p"):
        return None
    return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()


def _strip_section(doc: Document, start_text: str, end_text: str):
    """
    Remove the template's placeholder section running from the heading whose text
    is `start_text` up to (but not including) the heading `end_text`, and return
    the `end_text` element so the caller can insert real content in its place.

    For Exp 6 this deletes the blank "Observations" heading and its two example
    tables, so only the measured "Result" section is left — otherwise the empty
    sample tables sit right above the real ones. Returns None if either heading is
    missing (caller then falls back to appending after the last table).
    """
    body = doc.element.body
    children = list(body.iterchildren())
    start = end = None
    for i, el in enumerate(children):
        txt = _para_text(el)
        if txt == start_text and start is None:
            start = i
        elif txt == end_text:
            end = i
            break
    if start is None or end is None or end <= start:
        return None
    end_el = children[end]
    for el in children[start:end]:
        body.remove(el)
    return end_el


@router.post("/export/bluetooth")
async def export_bluetooth_document(
    devices: Optional[str] = Form(None),
    connected: Optional[str] = Form(None),
    readings: Optional[str] = Form(None),
    heading: str = Form("Result"),
    template: str = Form("exp6"),
):
    """
    Inserts a measured "Result" section below the range-test observation table of
    the Experiment 6 document: Table 1 (discovered devices), an optional
    connected-device table, and Table 2 (paced RSSI field test), plus an
    observation line. No graph — the syllabus records observations as tables only.
    """
    def _parse_list(name, raw):
        if not raw:
            return []
        try:
            obj = json.loads(raw)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail=f"'{name}' payload was not valid JSON.")
        return obj if isinstance(obj, list) else []

    device_list = _parse_list("devices", devices)
    reading_list = _parse_list("readings", readings)
    conn_data = None
    if connected:
        try:
            obj = json.loads(connected)
            conn_data = obj if isinstance(obj, dict) and obj else None
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="'connected' payload was not valid JSON.")
    if not device_list and not reading_list:
        raise HTTPException(status_code=400, detail="No Bluetooth results to export. Scan for devices and log some range readings first.")

    raw = _load_template(template)
    base = _TEMPLATES[template][:-5]  # "Expt. No. 6"

    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        raise HTTPException(status_code=400, detail="Could not open the experiment document.")

    anchor_table = _find_exp6_anchor(doc)
    if anchor_table is None:
        raise HTTPException(status_code=422, detail="No observation table found in the document.")
    try:
        style_hint = anchor_table.style
    except Exception:
        style_hint = None

    # Delete the template's blank "Observations" section (heading + the two example
    # tables); the measured Result section replaces it, right before Conclusion.
    insert_before = _strip_section(doc, "Observations", "Conclusion")

    built = [doc.add_paragraph()]  # blank before heading

    h = doc.add_paragraph()
    hr = h.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    built.append(h)
    built.append(doc.add_paragraph())

    if device_list:
        built.append(_sublabel(doc, "Table 1: Bluetooth Discovery & Service Profiling — Measured"))
        built.append(_build_bt_discovery_table(doc, device_list, style_hint=style_hint))
        built.append(doc.add_paragraph())

    if conn_data:
        built.append(_sublabel(doc, "Table 2: Connected / Paired Device — Service Mapping"))
        built.append(_build_bt_connected_table(doc, conn_data, style_hint=style_hint))
        built.append(doc.add_paragraph())

    if reading_list:
        n = 3 if conn_data else 2
        built.append(_sublabel(doc, f"Table {n}: Signal Quality vs. Distance Field Test — Measured"))
        built.append(_build_bt_field_table(doc, reading_list, style_hint=style_hint))
        built.append(doc.add_paragraph())

    obs = doc.add_paragraph()
    obs.add_run("Observation: ").bold = True
    obs.add_run(_exp6_summary(device_list, reading_list))
    built.append(obs)
    built.append(doc.add_paragraph())

    if insert_before is not None:
        # Place the Result section where the placeholder Observations section was.
        for block in built:
            insert_before.addprevious(block._element)
    else:
        # Fallback: no Observations/Conclusion headings found — append after the
        # last observation table, as before.
        anchor = anchor_table._element
        for block in built:
            el = block._element
            anchor.addnext(el)
            anchor = el

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    filename = f"{base} - with Results.docx"
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/export/network")
async def export_network_document(
    ping: Optional[str] = Form(None),
    speedtest: Optional[str] = Form(None),
    chart: Optional[UploadFile] = File(None),
    heading: str = Form("Result"),
    platform: str = Form("PC / Laptop (Browser)"),
    template: str = Form("exp5"),
):
    """
    Inserts a measured "Result" section below the throughput observation table of
    the Experiment 5 document and returns it. Accepts the real command-line ping
    result and/or the live speed-test result (at least one is required); builds a
    result table for whichever is present.
    """
    def _parse(name, raw):
        if not raw:
            return None
        try:
            obj = json.loads(raw)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail=f"'{name}' payload was not valid JSON.")
        return obj if isinstance(obj, dict) and obj else None

    ping_data = _parse("ping", ping)
    speed_data = _parse("speedtest", speedtest)
    if not ping_data and not speed_data:
        raise HTTPException(status_code=400, detail="No measured results to export. Run the ping and/or speed test first.")

    raw = _load_template(template)
    base = _TEMPLATES[template][:-5]  # e.g. "Expt No. 5"

    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        raise HTTPException(status_code=400, detail="Could not open the experiment document.")

    anchor_table = _find_exp5_anchor(doc)
    if anchor_table is None:
        raise HTTPException(
            status_code=422,
            detail="No observation table found in the document, so there is nowhere to insert the results.",
        )
    try:
        style_hint = anchor_table.style
    except Exception:
        style_hint = None

    built = []
    built.append(doc.add_paragraph())  # blank before heading

    h = doc.add_paragraph()
    hr = h.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    built.append(h)

    built.append(doc.add_paragraph())  # blank after heading

    if ping_data:
        built.append(_sublabel(doc, "Table 1: Command Line Interface (CLI) Results — Measured"))
        built.append(_build_cli_table(doc, ping_data, style_hint=style_hint))
        if ping_data.get("jitter") is not None:
            note = doc.add_paragraph()
            nr = note.add_run(f"Jitter (avg RTT variation): {_num(ping_data.get('jitter'))} ms")
            nr.font.size = Pt(9)
            nr.italic = True
            built.append(note)
        built.append(doc.add_paragraph())  # spacer

    if speed_data:
        built.append(_sublabel(doc, "Table 2: Online Simulation Results — Measured"))
        built.append(_build_throughput_table(doc, speed_data, platform, style_hint=style_hint))
        srv = (speed_data.get("server_name") or "").strip()
        if srv:
            note = doc.add_paragraph()
            nr = note.add_run(f"Test server: {srv}"
                              + (f", {speed_data.get('server_country')}" if speed_data.get("server_country") else ""))
            nr.font.size = Pt(9)
            nr.italic = True
            built.append(note)
        built.append(doc.add_paragraph())  # spacer

    if chart is not None:
        img_bytes = await chart.read()
        if img_bytes:
            cap = doc.add_paragraph()
            cr = cap.add_run("Graph: Throughput over time")
            cr.bold = True
            cr.font.size = Pt(10)
            built.append(cap)
            built.append(doc.add_paragraph())
            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                pic_para.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
            except Exception:
                pic_para.add_run("[Chart image could not be embedded]").italic = True
            built.append(pic_para)
            built.append(doc.add_paragraph())

    obs = doc.add_paragraph()
    obs.add_run("Observation: ").bold = True
    obs.add_run(_exp5_summary(ping_data, speed_data))
    built.append(obs)

    built.append(doc.add_paragraph())  # trailing blank before Conclusion

    anchor = anchor_table._element
    for block in built:
        el = block._element
        anchor.addnext(el)
        anchor = el

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    filename = f"{base} - with Results.docx"
    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/export")
async def export_to_document(
    readings: str = Form(...),
    document: Optional[UploadFile] = File(None),
    chart: Optional[UploadFile] = File(None),
    heading: str = Form("Result & Graph"),
    experiment: str = Form("Experiment 4 - Wi-Fi Signal Strength vs Distance"),
    template: str = Form("exp4"),
):
    """
    Inserts a "Result & Graph" section below the Observation Table of the
    experiment document and returns it. By default the bundled template for the
    given experiment is used (no upload); an uploaded .docx overrides it.
    """
    try:
        rows = json.loads(readings)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Readings payload was not valid JSON.")
    if not isinstance(rows, list) or not rows:
        raise HTTPException(status_code=400, detail="No readings to export.")

    # Source bytes + the base filename for the download. Prefer an uploaded file
    # (back-compat); otherwise fall back to the bundled template.
    if document is not None and document.filename:
        if not document.filename.lower().endswith(".docx"):
            raise HTTPException(status_code=400, detail="Please upload a .docx file (not .doc or .pdf).")
        raw = await document.read()
        base = document.filename[:-5]
    else:
        raw = _load_template(template)
        base = _TEMPLATES[template][:-5]  # e.g. "Expt. No. 4"

    try:
        doc = Document(io.BytesIO(raw))
    except Exception:
        raise HTTPException(
            status_code=400,
            detail="Could not open the experiment document. Make sure it is a valid .docx.",
        )

    anchor_table = _find_observation_table(doc)
    if anchor_table is None:
        raise HTTPException(
            status_code=422,
            detail="No Observation Table found in the document, so there is nowhere to insert the results.",
        )

    # --- Build the new content at the end of the body, then relocate it ---
    built = []

    built.append(doc.add_paragraph())  # blank line before the heading

    h = doc.add_paragraph()
    hr = h.add_run(heading)
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    built.append(h)

    built.append(doc.add_paragraph())  # blank line after the heading

    try:
        style_hint = anchor_table.style
    except Exception:
        style_hint = None
    tbl = _build_readings_table(doc, rows, style_hint=style_hint)
    built.append(tbl)

    # Space between the readings table and the graph (matches the syllabus doc).
    for _ in range(_GRAPH_GAP):
        built.append(doc.add_paragraph())

    if chart is not None:
        img_bytes = await chart.read()
        if img_bytes:
            cap = doc.add_paragraph()
            cr = cap.add_run("Graph: Signal Strength vs Distance")
            cr.bold = True
            cr.font.size = Pt(10)
            built.append(cap)

            built.append(doc.add_paragraph())  # blank line before the image

            pic_para = doc.add_paragraph()
            pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                pic_para.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(6.4))
            except Exception:
                # A bad/empty image must not lose the table the student just measured.
                pic_para.add_run("[Chart image could not be embedded]").italic = True
            built.append(pic_para)

            built.append(doc.add_paragraph())  # blank line after the image

    obs = doc.add_paragraph()
    obs.add_run("Observation: ").bold = True
    obs.add_run(_summary_sentence(rows))
    built.append(obs)

    built.append(doc.add_paragraph())  # trailing blank line before Conclusion

    # --- Move the built blocks to sit directly beneath the Observation Table ---
    anchor = anchor_table._element
    for block in built:
        el = block._element
        anchor.addnext(el)
        anchor = el

    # The graph now lives in the document, so drop the "Paste the Graph"
    # placeholder that the template used for a hand-pasted image.
    _remove_graph_placeholder(doc)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)

    filename = f"{base} - with Results.docx"

    return StreamingResponse(
        out,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
